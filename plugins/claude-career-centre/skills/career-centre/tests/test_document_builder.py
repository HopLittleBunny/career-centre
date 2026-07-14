from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

FIXTURE_ROOT = Path(__file__).resolve().parent / "fixtures"
sys.path.insert(0, str(FIXTURE_ROOT))

from factory import application_pack, earlycareer_application_pack  # noqa: E402


class DocumentBuilderTests(unittest.TestCase):
    def test_builder_creates_paired_structurally_valid_docx_files(self) -> None:
        skill_root = Path(__file__).resolve().parents[1]
        builder = skill_root / "scripts" / "build_application_pack.py"
        with tempfile.TemporaryDirectory() as temporary:
            temporary_path = Path(temporary)
            input_path = temporary_path / "input.json"
            output_path = temporary_path / "pack"
            input_path.write_text(json.dumps(application_pack(), indent=2), encoding="utf-8")
            process = subprocess.run(
                [sys.executable, str(builder), "--input", str(input_path), "--output-dir", str(output_path)],
                capture_output=True,
                text=True,
            )
            self.assertEqual(process.returncode, 0, process.stdout + process.stderr)
            manifest = json.loads((output_path / "application_pack_manifest.json").read_text(encoding="utf-8"))
            self.assertEqual(manifest["status"], "BUILT_STRUCTURAL_PASS")
            self.assertTrue((output_path / manifest["cv"]).exists())
            self.assertTrue((output_path / manifest["cover_letter"]).exists())
            self.assertTrue(all(not Path(item["path"]).is_absolute() for item in manifest["files"]))
            for report_name in manifest["structural_reports"]:
                report = json.loads((output_path / report_name).read_text(encoding="utf-8"))
                self.assertTrue(report["passed"], report)

    def test_cv_only_is_explicit_and_does_not_create_cover_letter(self) -> None:
        skill_root = Path(__file__).resolve().parents[1]
        builder = skill_root / "scripts" / "build_application_pack.py"
        with tempfile.TemporaryDirectory() as temporary:
            temporary_path = Path(temporary)
            input_path = temporary_path / "input.json"
            output_path = temporary_path / "pack"
            input_path.write_text(json.dumps(application_pack(cv_only=True), indent=2), encoding="utf-8")
            process = subprocess.run(
                [sys.executable, str(builder), "--input", str(input_path), "--output-dir", str(output_path)],
                capture_output=True,
                text=True,
            )
            self.assertEqual(process.returncode, 0, process.stdout + process.stderr)
            manifest = json.loads((output_path / "application_pack_manifest.json").read_text(encoding="utf-8"))
            self.assertTrue(manifest["cv_only"])
            self.assertIsNone(manifest["cover_letter"])

    def test_reference_cv_can_supply_the_format_model(self) -> None:
        skill_root = Path(__file__).resolve().parents[1]
        builder = skill_root / "scripts" / "build_application_pack.py"
        reference = skill_root / "assets" / "templates" / "Professional_ATS_CV_Template.docx"
        data = application_pack()
        data["document_settings"]["template"] = "reference"
        data["document_settings"]["reference_template_name"] = "Reference_CV.docx"
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            input_path = root / "input.json"
            output_path = root / "pack"
            input_path.write_text(json.dumps(data, indent=2), encoding="utf-8")
            process = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    "--input",
                    str(input_path),
                    "--output-dir",
                    str(output_path),
                    "--cv-template",
                    str(reference),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(process.returncode, 0, process.stdout + process.stderr)
            manifest = json.loads((output_path / "application_pack_manifest.json").read_text(encoding="utf-8"))
            change_log = (output_path / "Change_Log.md").read_text(encoding="utf-8")
        self.assertEqual(manifest["format_source"], "Reference_CV.docx")
        self.assertIn("Reference_CV.docx", change_log)

    def test_reference_format_is_distilled_without_reference_content(self) -> None:
        skill_root = Path(__file__).resolve().parents[1]
        builder = skill_root / "scripts" / "build_application_pack.py"
        data = application_pack()
        data["document_settings"]["template"] = "reference"
        data["document_settings"]["reference_template_name"] = "Synthetic_Georgia_Reference.docx"
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            reference = root / "reference.docx"
            source = Document()
            source.styles["Normal"].font.name = "Georgia"
            source.styles["Normal"].font.size = Pt(10.5)
            source.styles["Heading 1"].font.name = "Georgia"
            source.styles["Heading 1"].font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
            source.styles["Title"].font.name = "Georgia"
            source.styles["Title"].font.color.rgb = RGBColor(0x13, 0x4E, 0x4A)
            section = source.sections[0]
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            source.add_paragraph("REFERENCE PERSON PRIVATE CONTENT")
            source.sections[0].header.add_paragraph("PRIVATE REFERENCE HEADER")
            source.core_properties.author = "Reference Person"
            source.save(reference)
            input_path = root / "input.json"
            output_path = root / "pack"
            input_path.write_text(json.dumps(data, indent=2), encoding="utf-8")
            process = subprocess.run(
                [
                    sys.executable,
                    str(builder),
                    "--input",
                    str(input_path),
                    "--output-dir",
                    str(output_path),
                    "--cv-template",
                    str(reference),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(process.returncode, 0, process.stdout + process.stderr)
            manifest = json.loads((output_path / "application_pack_manifest.json").read_text(encoding="utf-8"))
            output_doc = Document(output_path / manifest["cv"])
            visible_text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
            header_text = "\n".join(paragraph.text for paragraph in output_doc.sections[0].header.paragraphs)
        self.assertEqual(output_doc.styles["Normal"].font.name, "Georgia")
        self.assertEqual(str(output_doc.styles["CCC Heading"].font.color.rgb), "0F766E")
        self.assertAlmostEqual(output_doc.sections[0].left_margin.inches, 0.8, places=1)
        self.assertNotIn("REFERENCE PERSON PRIVATE CONTENT", visible_text)
        self.assertNotIn("PRIVATE REFERENCE HEADER", header_text)
        self.assertEqual(output_doc.core_properties.author, "")

    def test_reference_mode_fails_without_the_reference_docx(self) -> None:
        skill_root = Path(__file__).resolve().parents[1]
        builder = skill_root / "scripts" / "build_application_pack.py"
        data = application_pack()
        data["document_settings"]["template"] = "reference"
        data["document_settings"]["reference_template_name"] = "Reference_CV.docx"
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            input_path = root / "input.json"
            input_path.write_text(json.dumps(data, indent=2), encoding="utf-8")
            process = subprocess.run(
                [sys.executable, str(builder), "--input", str(input_path), "--output-dir", str(root / "pack")],
                capture_output=True,
                text=True,
            )
        self.assertNotEqual(process.returncode, 0)
        self.assertIn("requires --cv-template", process.stdout)

    def test_earlycareer_pack_builds_as_one_page_target(self) -> None:
        skill_root = Path(__file__).resolve().parents[1]
        builder = skill_root / "scripts" / "build_application_pack.py"
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            input_path = root / "input.json"
            output_path = root / "pack"
            input_path.write_text(json.dumps(earlycareer_application_pack(), indent=2), encoding="utf-8")
            process = subprocess.run(
                [sys.executable, str(builder), "--input", str(input_path), "--output-dir", str(output_path)],
                capture_output=True,
                text=True,
            )
            self.assertEqual(process.returncode, 0, process.stdout + process.stderr)
            report = json.loads((output_path / "CV_Structural_Validation.json").read_text(encoding="utf-8"))
        self.assertEqual(report["metrics"]["manual_page_breaks"], 0)

    def test_validator_rejects_private_use_bullets_and_generic_cover_opening(self) -> None:
        skill_root = Path(__file__).resolve().parents[1]
        validator = skill_root / "scripts" / "validate_docx.py"
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            broken_cv = root / "Broken_CV.docx"
            cv = Document()
            cv.add_paragraph("This built-in bullet may render as a hollow box.", style="List Bullet")
            cv.save(broken_cv)
            cv_result = subprocess.run(
                [sys.executable, str(validator), str(broken_cv)],
                capture_output=True,
                text=True,
            )

            broken_cover = root / "Broken_Cover_Letter.docx"
            cover = Document()
            cover.add_paragraph(
                "I am applying for this role because my experience appears to align with the opportunity and I would like to be considered by the hiring team."
            )
            cover.save(broken_cover)
            cover_result = subprocess.run(
                [sys.executable, str(validator), str(broken_cover)],
                capture_output=True,
                text=True,
            )

        self.assertNotEqual(cv_result.returncode, 0)
        self.assertIn("Private-use glyphs detected", cv_result.stdout)
        self.assertNotEqual(cover_result.returncode, 0)
        self.assertIn("Generic cover-letter opening detected", cover_result.stdout)

    def test_validator_accepts_portable_unicode_bullet(self) -> None:
        skill_root = Path(__file__).resolve().parents[1]
        validator = skill_root / "scripts" / "validate_docx.py"
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            portable_cv = root / "Portable_CV.docx"
            cv = Document()
            paragraph = cv.add_paragraph()
            run = paragraph.add_run("\u2022 Portable bullet in the normal body font.")
            run.font.name = "Arial"
            cv.core_properties.author = ""
            cv.core_properties.last_modified_by = ""
            cv.core_properties.comments = ""
            cv.core_properties.keywords = ""
            cv.save(portable_cv)
            result = subprocess.run(
                [sys.executable, str(validator), str(portable_cv)],
                capture_output=True,
                text=True,
            )
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        report = json.loads(result.stdout)
        self.assertEqual(report["metrics"]["private_use_glyphs"], 0)
        self.assertEqual(report["metrics"]["active_numbering_private_use_glyphs"], 0)


if __name__ == "__main__":
    unittest.main()
