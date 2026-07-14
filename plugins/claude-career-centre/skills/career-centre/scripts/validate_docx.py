#!/usr/bin/env python3
"""Structural, content and safety validation for generated DOCX files."""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document

URL_RE = re.compile(r"(?i)\b(?:https?://|www\.)\S+")
PRIVATE_USE_GLYPH_RE = re.compile(r"[\uE000-\uF8FF]")
XML_ENTITY_RE = re.compile(r"&#(?:x([0-9A-Fa-f]+)|([0-9]+));")
RISKY_BULLET_FONTS = {"wingdings", "wingdings 2", "wingdings 3", "symbol"}
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": WORD_NS}
GENERIC_COVER_OPENERS = (
    "i am applying",
    "i'm applying",
    "i’m applying",
    "i am excited to apply",
    "i'm excited to apply",
    "i’m excited to apply",
    "please accept my application",
)
PLACEHOLDERS = (
    "candidate name", "replace with", "lorem ipsum", "todo", "[[", "]]",
    "target positioning line", "skill 1", "role title",
)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _effective_size(run: Any, paragraph: Any, doc: Document) -> float | None:
    if run.font.size:
        return run.font.size.pt
    try:
        if paragraph.style and paragraph.style.font.size:
            return paragraph.style.font.size.pt
    except Exception:
        pass
    try:
        if doc.styles["Normal"].font.size:
            return doc.styles["Normal"].font.size.pt
    except Exception:
        pass
    return None


def _private_use_count(value: str) -> int:
    count = len(PRIVATE_USE_GLYPH_RE.findall(value))
    for match in XML_ENTITY_RE.finditer(value):
        codepoint = int(match.group(1), 16) if match.group(1) else int(match.group(2))
        if 0xE000 <= codepoint <= 0xF8FF:
            count += 1
    return count


def _active_numbering_risks(document_xml: bytes, styles_xml: bytes | None, numbering_xml: bytes | None) -> tuple[int, list[str]]:
    """Return private-use glyph count and legacy fonts for numbering actually used.

    DOCX packages often contain unused built-in Wingdings numbering definitions.
    Rejecting the whole numbering part would therefore create false positives.
    This follows paragraph numIds and paragraph styles used by the document to
    the active abstract numbering definitions only.
    """
    if not numbering_xml:
        return 0, []
    document_root = ET.fromstring(document_xml)
    used_num_ids = {
        node.get(f"{{{WORD_NS}}}val")
        for node in document_root.findall(".//w:pPr/w:numPr/w:numId", NS)
        if node.get(f"{{{WORD_NS}}}val") not in {None, "0"}
    }
    used_style_ids = {
        node.get(f"{{{WORD_NS}}}val")
        for node in document_root.findall(".//w:pPr/w:pStyle", NS)
        if node.get(f"{{{WORD_NS}}}val")
    }
    if styles_xml and used_style_ids:
        styles_root = ET.fromstring(styles_xml)
        styles_by_id = {
            style.get(f"{{{WORD_NS}}}styleId"): style
            for style in styles_root.findall("w:style", NS)
        }
        pending = list(used_style_ids)
        visited: set[str] = set()
        while pending:
            style_id = pending.pop()
            if style_id in visited:
                continue
            visited.add(style_id)
            style = styles_by_id.get(style_id)
            if style is None:
                continue
            num_id = style.find("./w:pPr/w:numPr/w:numId", NS)
            if num_id is not None:
                value = num_id.get(f"{{{WORD_NS}}}val")
                if value and value != "0":
                    used_num_ids.add(value)
            based_on = style.find("./w:basedOn", NS)
            if based_on is not None:
                parent_id = based_on.get(f"{{{WORD_NS}}}val")
                if parent_id:
                    pending.append(parent_id)

    numbering_root = ET.fromstring(numbering_xml)
    active_abstract_ids: set[str] = set()
    active_num_nodes: list[ET.Element] = []
    for num in numbering_root.findall("w:num", NS):
        num_id = num.get(f"{{{WORD_NS}}}numId")
        if num_id not in used_num_ids:
            continue
        active_num_nodes.append(num)
        abstract = num.find("w:abstractNumId", NS)
        if abstract is not None:
            value = abstract.get(f"{{{WORD_NS}}}val")
            if value:
                active_abstract_ids.add(value)

    active_levels: list[ET.Element] = []
    for abstract in numbering_root.findall("w:abstractNum", NS):
        if abstract.get(f"{{{WORD_NS}}}abstractNumId") in active_abstract_ids:
            active_levels.extend(abstract.findall("w:lvl", NS))
    for num in active_num_nodes:
        active_levels.extend(num.findall(".//w:lvl", NS))

    private_use = 0
    risky_fonts: set[str] = set()
    for level in active_levels:
        for label in level.findall(".//w:lvlText", NS):
            private_use += _private_use_count(label.get(f"{{{WORD_NS}}}val", ""))
        for fonts in level.findall(".//w:rFonts", NS):
            for attribute in ("ascii", "hAnsi", "cs", "eastAsia"):
                value = fonts.get(f"{{{WORD_NS}}}{attribute}")
                if value and value.casefold() in RISKY_BULLET_FONTS:
                    risky_fonts.add(value)
    return private_use, sorted(risky_fonts, key=str.casefold)


def validate(path: Path, rules: dict[str, Any] | None = None) -> dict[str, Any]:
    rules = rules or {}
    errors: list[str] = []
    warnings: list[str] = []
    report: dict[str, Any] = {
        "file": str(path),
        "errors": errors,
        "warnings": warnings,
        "metrics": {},
        "hyperlinks": [],
    }
    if not path.exists():
        errors.append("DOCX file does not exist.")
        report["passed"] = False
        return report
    report["sha256"] = _sha256(path)
    document_type = rules.get("document_type") or ("cover_letter" if "cover_letter" in path.stem.casefold() else "cv")
    report["metrics"]["document_type"] = document_type
    try:
        doc = Document(path)
    except Exception as exc:
        errors.append(f"DOCX does not open: {exc}")
        report["passed"] = False
        return report

    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    header_footer_text: list[str] = []
    for section in doc.sections:
        for part in (
            section.header, section.footer, section.first_page_header,
            section.first_page_footer, section.even_page_header,
            section.even_page_footer,
        ):
            header_footer_text.extend(paragraph.text for paragraph in part.paragraphs)
    text = "\n".join([body_text, *header_footer_text])
    lower = text.casefold()
    report["metrics"].update(
        {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "words": len(text.split()),
            "body_words": len(body_text.split()),
            "sections": len(doc.sections),
        }
    )
    if doc.tables and not rules.get("allow_tables", False):
        errors.append("Tables detected; ATS-safe output must not use layout tables.")
    raw_urls = URL_RE.findall(text)
    if raw_urls:
        errors.append(f"Visible raw URLs detected: {raw_urls[:5]}")
    minimum_font = float(rules.get("minimum_font_pt", 9))
    small_fonts: list[float] = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            size = _effective_size(run, paragraph, doc)
            if size is not None and size < minimum_font - 0.01:
                small_fonts.append(size)
    if small_fonts:
        errors.append(f"Font below {minimum_font:g} pt detected; minimum effective size {min(small_fonts):.1f} pt.")

    for placeholder in rules.get("placeholders", PLACEHOLDERS):
        if str(placeholder).casefold() in lower:
            errors.append(f"Unresolved placeholder detected: {placeholder}")
    for term in rules.get("banned_terms", []):
        if str(term).casefold() in lower:
            errors.append(f"Banned or unsafe term detected: {term}")
    for required in rules.get("required_text", []):
        if str(required).casefold() not in lower:
            errors.append(f"Required text missing: {required}")

    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        document_xml_bytes = archive.read("word/document.xml")
        xml = document_xml_bytes.decode("utf-8", "ignore")
        visible_xml_names = [
            name
            for name in names
            if name == "word/document.xml"
            or re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
            or name in {"word/footnotes.xml", "word/endnotes.xml"}
        ]
        visible_private_use = sum(
            _private_use_count(archive.read(name).decode("utf-8", "ignore"))
            for name in visible_xml_names
        )
        numbering_private_use, risky_numbering_fonts = _active_numbering_risks(
            document_xml_bytes,
            archive.read("word/styles.xml") if "word/styles.xml" in names else None,
            archive.read("word/numbering.xml") if "word/numbering.xml" in names else None,
        )
        private_use_count = visible_private_use + numbering_private_use
        report["metrics"]["private_use_glyphs"] = private_use_count
        report["metrics"]["active_numbering_private_use_glyphs"] = numbering_private_use
        report["metrics"]["active_numbering_legacy_fonts"] = risky_numbering_fonts
        if private_use_count:
            errors.append(
                "Private-use glyphs detected in visible content or active Word numbering; these commonly render "
                "as hollow boxes. Use U+2022 in the normal body font, not Wingdings/Symbol numbering."
            )
        if risky_numbering_fonts:
            errors.append(
                "Legacy bullet font detected in active Word numbering: "
                + ", ".join(risky_numbering_fonts)
                + ". Use U+2022 in the normal body font."
            )
        if "w:txbxContent" in xml:
            errors.append("Text-box content detected.")
        if "w:commentRangeStart" in xml or "word/comments.xml" in names:
            errors.append("Comments detected.")
        if "<w:ins" in xml or "<w:del" in xml:
            errors.append("Tracked changes detected.")
        if "w:vanish" in xml:
            errors.append("Hidden text detected.")
        media = [name for name in names if name.startswith("word/media/")]
        if media and not rules.get("allow_media", False):
            errors.append(f"Embedded media detected: {len(media)} file(s).")
        relationship_name = "word/_rels/document.xml.rels"
        if relationship_name in names:
            rel_xml = archive.read(relationship_name).decode("utf-8", "ignore")
            targets = re.findall(r'Type="[^"]*/hyperlink"[^>]*Target="([^"]+)"', rel_xml)
            report["hyperlinks"] = targets
            report["metrics"]["hyperlinks"] = len(targets)
            unsafe = [target for target in targets if not re.match(r"(?i)^(?:https?://|mailto:)", target)]
            if unsafe:
                errors.append(f"Unsafe hyperlink target(s) detected: {unsafe[:5]}")
        else:
            report["metrics"]["hyperlinks"] = 0
        manual_breaks = xml.count('w:type="page"')
        report["metrics"]["manual_page_breaks"] = manual_breaks
        expected_breaks = rules.get("expected_manual_page_breaks")
        if expected_breaks is not None and manual_breaks != int(expected_breaks):
            errors.append(f"Expected {expected_breaks} controlled page break(s); found {manual_breaks}.")

    core = doc.core_properties
    if core.author or core.last_modified_by or core.comments or core.keywords:
        errors.append("Personal or internal document metadata is populated.")
    word_count = report["metrics"]["body_words"]
    if document_type == "cover_letter":
        narrative_paragraph = next(
            (
                paragraph.text.strip()
                for paragraph in doc.paragraphs
                if len(paragraph.text.strip().split()) >= 12
            ),
            "",
        )
        opening = narrative_paragraph.casefold()
        if opening.startswith(GENERIC_COVER_OPENERS):
            errors.append(
                "Generic cover-letter opening detected. Start with a natural, role-specific reason or thesis, "
                "not 'I am applying' or 'I am excited to apply'."
            )
        minimum = int(rules.get("minimum_words", 220))
        maximum = int(rules.get("maximum_words", 550))
        if word_count < minimum:
            warnings.append(f"Cover letter is short at {word_count} words; inspect narrative depth.")
        if word_count > maximum:
            errors.append(f"Cover letter is too long at {word_count} words; maximum is {maximum}.")
    else:
        minimum = int(rules.get("minimum_words", 300))
        maximum = int(rules.get("maximum_words", 1450))
        if word_count < minimum:
            warnings.append(f"CV is sparse at {word_count} words; inspect role evidence and page density.")
        if word_count > maximum:
            warnings.append(f"CV is dense at {word_count} words; inspect readability and page count.")
    if not text.strip():
        errors.append("Document contains no visible text.")
    report["passed"] = not errors
    return report


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx")
    parser.add_argument("--rules")
    parser.add_argument("--json-out")
    args = parser.parse_args()
    rules = json.loads(Path(args.rules).read_text(encoding="utf-8")) if args.rules else {}
    report = validate(Path(args.docx), rules)
    payload = json.dumps(report, indent=2)
    if args.json_out:
        Path(args.json_out).write_text(payload, encoding="utf-8")
    print(payload)
    return 0 if report["passed"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
