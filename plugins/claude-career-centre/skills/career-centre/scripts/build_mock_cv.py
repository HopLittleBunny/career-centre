#!/usr/bin/env python3
"""Create a polished synthetic source CV for live-plugin evaluation."""
from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor

NAVY = "17324D"
GREY = "4B5563"
LIGHT_GREY = "D6DEE6"


def _font(style, size: float, *, bold: bool = False, color: str | None = None) -> None:
    style.font.name = "Arial"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def _hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), NAVY)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    properties.extend([fonts, colour, underline, size])
    run.append(properties)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Mock Heading")
    paragraph.add_run(text.upper())
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), LIGHT_GREY)
    borders.append(bottom)
    properties.append(borders)


def build(persona: dict, output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)
    styles = doc.styles
    _font(styles["Normal"], 10.2)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Normal"].paragraph_format.line_spacing = 1.05
    for name, size, bold, color, before, after in (
        ("Mock Name", 21, True, NAVY, 0, 1),
        ("Mock Headline", 10.8, True, GREY, 0, 2),
        ("Mock Contact", 9.2, False, GREY, 0, 7),
        ("Mock Heading", 10.5, True, NAVY, 7, 3),
        ("Mock Experience", 10.2, True, NAVY, 3, 1),
    ):
        style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        _font(style, size, bold=bold, color=color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    bullet = styles["List Bullet"]
    _font(bullet, 10.0)
    bullet.paragraph_format.left_indent = Inches(0.23)
    bullet.paragraph_format.first_line_indent = Inches(-0.16)
    bullet.paragraph_format.space_after = Pt(2.5)
    bullet.paragraph_format.keep_together = True

    profile = persona["profile"]
    source = persona["source_cv"]
    paragraph = doc.add_paragraph(style="Mock Name")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(profile["name"])
    paragraph = doc.add_paragraph(style="Mock Headline")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(profile["headline"])
    paragraph = doc.add_paragraph(style="Mock Contact")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    values = [{"text": profile["location"], "url": None}, *profile.get("contact", [])]
    for index, item in enumerate(values):
        if index:
            paragraph.add_run("  |  ")
        text = item.get("text", "")
        url = item.get("url")
        if url:
            label = "Email" if str(url).startswith("mailto:") else text
            _hyperlink(paragraph, label, url)
        else:
            paragraph.add_run(text)

    _heading(doc, "Professional Summary")
    doc.add_paragraph(source["summary"])
    _heading(doc, "Selected Experience")
    for experience in source["experience"]:
        paragraph = doc.add_paragraph(style="Mock Experience")
        pieces = [experience.get(key, "") for key in ("role", "company", "location", "dates") if experience.get(key)]
        paragraph.add_run(" | ".join(pieces))
        for bullet_text in experience.get("bullets", []):
            doc.add_paragraph(bullet_text, style="List Bullet")
    _heading(doc, "Core Skills")
    doc.add_paragraph(" | ".join(source["skills"]))
    _heading(doc, "Education and Certification")
    for item in source["education"]:
        doc.add_paragraph(item)
    properties = doc.core_properties
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""
    properties.title = ""
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--persona", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()
    persona = json.loads(Path(args.persona).read_text(encoding="utf-8"))
    build(persona, Path(args.output))
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
