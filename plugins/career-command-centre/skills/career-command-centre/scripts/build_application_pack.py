#!/usr/bin/env python3
"""Build an evidence-mapped Word CV and paired cover letter from v4 JSON."""
from __future__ import annotations

import argparse
import hashlib
import json
import math
import re
import shutil
import unicodedata
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor

from contracts import validate_application_pack
from validate_docx import validate as validate_docx

NAVY = "17324D"
GREY = "4B5563"
LIGHT_GREY = "D6DEE6"


def _safe_name(value: str) -> str:
    normalised = unicodedata.normalize("NFKC", str(value).strip())
    cleaned = "".join(character if (character.isalnum() or character in {"_", "-"}) else "_" for character in normalised)
    return re.sub(r"_+", "_", cleaned).strip("_-") or "Document"


def _sha256(path: Path) -> str:
    digest = hashlib.sha256(path.read_bytes())
    return digest.hexdigest()


def _clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _clear_headers_and_footers(doc: Document) -> None:
    for section in doc.sections:
        for part in (
            section.header, section.footer, section.first_page_header,
            section.first_page_footer, section.even_page_header,
            section.even_page_footer,
        ):
            for paragraph in list(part.paragraphs):
                paragraph._element.getparent().remove(paragraph._element)
            for table in list(part.tables):
                table._element.getparent().remove(table._element)


def _set_font(run: Any, name: str, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    if run._element.get_or_add_rPr().rFonts is None:
        fonts = OxmlElement("w:rFonts")
        run._element.get_or_add_rPr().append(fonts)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_style_font(style: Any, name: str, size: float, bold: bool = False, color: str | None = None) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def _style_font_name(style: Any, fallback: str) -> str:
    return str(style.font.name).strip() if style is not None and style.font.name else fallback


def _style_font_size(style: Any, fallback: float) -> float:
    if style is not None and style.font.size is not None:
        return float(style.font.size.pt)
    return fallback


def _style_color(style: Any, fallback: str) -> str:
    if style is not None and style.font.color is not None and style.font.color.rgb is not None:
        return str(style.font.color.rgb)
    return fallback


def _new_doc(
    template: Path,
    *,
    page_size: str,
    minimum_font: float,
    cover: bool = False,
    reference_mode: bool = False,
) -> Document:
    doc = Document(template) if template.exists() else Document()
    styles = doc.styles
    reference_normal = styles["Normal"] if "Normal" in styles else None
    reference_heading = styles["Heading 1"] if "Heading 1" in styles else None
    reference_title = styles["Title"] if "Title" in styles else reference_heading
    body_font = _style_font_name(reference_normal, "Arial") if reference_mode else "Arial"
    heading_font = _style_font_name(reference_heading, body_font) if reference_mode else "Arial"
    title_font = _style_font_name(reference_title, heading_font) if reference_mode else "Arial"
    heading_color = _style_color(reference_heading, NAVY) if reference_mode else NAVY
    title_color = _style_color(reference_title, NAVY) if reference_mode else NAVY
    reference_margins = None
    if reference_mode and doc.sections:
        original = doc.sections[0]
        reference_margins = tuple(
            value.inches if value is not None else None
            for value in (original.top_margin, original.bottom_margin, original.left_margin, original.right_margin)
        )
    _clear_body(doc)
    _clear_headers_and_footers(doc)
    section = doc.sections[0]
    if page_size == "Letter":
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    else:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    margin = 0.72 if cover else 0.58
    if reference_margins:
        top, bottom, left, right = reference_margins
        section.top_margin = Inches(top if top is not None and 0.4 <= top <= 1.0 else margin)
        section.bottom_margin = Inches(bottom if bottom is not None and 0.4 <= bottom <= 1.0 else margin)
        section.left_margin = Inches(left if left is not None and 0.45 <= left <= 1.1 else margin + 0.08)
        section.right_margin = Inches(right if right is not None and 0.45 <= right <= 1.1 else margin + 0.08)
    else:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin + 0.08)
        section.right_margin = Inches(margin + 0.08)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    resolved_minimum_font = math.ceil(minimum_font * 2) / 2
    normal = styles["Normal"]
    reference_body_size = _style_font_size(reference_normal, 9.5) if reference_mode else 9.5
    body_size = max(resolved_minimum_font, min(reference_body_size, 11.5))
    _set_style_font(normal, body_font, body_size)
    normal.paragraph_format.space_after = Pt(2.3 if not cover else 7)
    normal.paragraph_format.line_spacing = 1.0 if not cover else 1.08

    def ensure(name: str, size: float, bold: bool, color: str, before: float, after: float, font_name: str = heading_font) -> Any:
        style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        _set_style_font(style, font_name, size, bold, color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        return style

    ensure("CCC Name", 20 if not cover else 17, True, title_color, 0, 1, title_font)
    ensure("CCC Headline", 10.5, True, GREY, 0, 2, body_font)
    ensure("CCC Contact", max(9.0, resolved_minimum_font), False, GREY, 0, 5, body_font)
    ensure("CCC Heading", 10.2, True, heading_color, 5, 2)
    ensure("CCC Experience", max(9.5, resolved_minimum_font), True, heading_color, 2, 0)

    # Do not use Word's built-in List Bullet style here. Its numbering definition
    # commonly stores U+F0B7 in Wingdings, which renders as a hollow square in
    # browser previews when that legacy font is unavailable. A literal U+2022 in
    # the normal body font is portable across Word, browser and PDF renderers.
    bullet = styles["CCC Bullet"] if "CCC Bullet" in styles else styles.add_style("CCC Bullet", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(bullet, body_font, max(body_size, resolved_minimum_font))
    bullet.paragraph_format.left_indent = Inches(0.2)
    bullet.paragraph_format.first_line_indent = Inches(-0.14)
    bullet.paragraph_format.space_after = Pt(1.7)
    bullet.paragraph_format.keep_together = True

    properties = doc.core_properties
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""
    properties.keywords = ""
    properties.subject = ""
    properties.title = ""
    return doc


def _add_border(paragraph: Any) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), LIGHT_GREY)
    borders.append(bottom)


def _add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), NAVY)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    run_properties.extend([fonts, colour, underline, size])
    run.append(run_properties)
    value = OxmlElement("w:t")
    value.text = str(text)
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _item_text(item: dict[str, Any]) -> str:
    return str(item.get("text", "")).strip()


def _add_header(doc: Document, candidate: dict[str, Any]) -> None:
    paragraph = doc.add_paragraph(style="CCC Name")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(candidate.get("name", ""))
    headline = str(candidate.get("headline", "")).strip()
    if headline:
        paragraph = doc.add_paragraph(style="CCC Headline")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(headline)
    contact = [item for item in candidate.get("contact", []) if isinstance(item, dict) and str(item.get("text", "")).strip()]
    if contact:
        paragraph = doc.add_paragraph(style="CCC Contact")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for index, item in enumerate(contact):
            if index:
                paragraph.add_run("  |  ")
            if item.get("url"):
                _add_hyperlink(paragraph, str(item["text"]), str(item["url"]))
            else:
                paragraph.add_run(str(item["text"]))


def _add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="CCC Bullet")
    paragraph.add_run("\u2022 ")
    paragraph.add_run(text)


def _build_cv(data: dict[str, Any], output_dir: Path, template: Path) -> Path:
    settings = data["document_settings"]
    doc = _new_doc(
        template,
        page_size=settings["page_size"],
        minimum_font=float(settings["minimum_font_pt"]),
        reference_mode=settings.get("template") == "reference",
    )
    _add_header(doc, data["candidate"])
    for section in data["cv"]["sections"]:
        if section["page_break_before"]:
            doc.add_page_break()
        heading = doc.add_paragraph(style="CCC Heading")
        heading.add_run(str(section["heading"]).upper())
        _add_border(heading)
        section_type = section["type"]
        if section_type == "paragraphs":
            for item in section["items"]:
                paragraph = doc.add_paragraph(_item_text(item))
                paragraph.paragraph_format.keep_together = True
        elif section_type == "bullets":
            for item in section["items"]:
                _add_bullet(doc, _item_text(item))
        elif section_type == "skills":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.keep_together = True
            paragraph.add_run(" | ".join(_item_text(item) for item in section["items"]))
        elif section_type == "experience":
            for experience in section["items"]:
                pieces = [
                    str(experience.get(key, "")).strip()
                    for key in ("role", "company", "location", "dates")
                    if str(experience.get(key, "")).strip()
                ]
                paragraph = doc.add_paragraph(style="CCC Experience")
                paragraph.add_run(" | ".join(pieces))
                for bullet in experience.get("bullets", []):
                    _add_bullet(doc, _item_text(bullet))
    identity = data["role"]["identity"]
    stem = "_".join(
        [
            _safe_name(data["candidate"]["name"]),
            _safe_name(identity["company"]),
            _safe_name(identity["title"]),
            "CV",
        ]
    )
    path = output_dir / f"{stem}.docx"
    doc.save(path)
    return path


def _build_cover_letter(data: dict[str, Any], output_dir: Path, template: Path) -> Path | None:
    cover = data["cover_letter"]
    if not cover["enabled"]:
        return None
    settings = data["document_settings"]
    doc = _new_doc(
        template,
        page_size=settings["page_size"],
        minimum_font=max(10, float(settings["minimum_font_pt"])),
        cover=True,
    )
    candidate = data["candidate"]
    identity = data["role"]["identity"]
    paragraph = doc.add_paragraph(style="CCC Name")
    paragraph.add_run(candidate["name"])
    contact = [item for item in candidate.get("contact", []) if str(item.get("text", "")).strip()]
    if contact:
        paragraph = doc.add_paragraph(style="CCC Contact")
        for index, item in enumerate(contact):
            if index:
                paragraph.add_run("  |  ")
            if item.get("url"):
                _add_hyperlink(paragraph, str(item["text"]), str(item["url"]))
            else:
                paragraph.add_run(str(item["text"]))
    for value in (cover.get("date"), cover.get("recipient"), identity.get("company")):
        if value:
            doc.add_paragraph(str(value))
    doc.add_paragraph(f"Re: {identity['title']}")
    doc.add_paragraph(cover["salutation"])
    for item in cover["paragraphs"]:
        doc.add_paragraph(_item_text(item))
    doc.add_paragraph(cover["closing"])
    doc.add_paragraph(cover["signature"])
    stem = "_".join(
        [
            _safe_name(candidate["name"]),
            _safe_name(identity["company"]),
            _safe_name(identity["title"]),
            "Cover_Letter",
        ]
    )
    path = output_dir / f"{stem}.docx"
    doc.save(path)
    return path


def _write_notes(data: dict[str, Any], output_dir: Path) -> list[Path]:
    change = data["change_log"]

    def bullets(values: list[str]) -> str:
        return "\n".join(f"- {value}" for value in values) or "- None"

    change_path = output_dir / "Change_Log.md"
    safety_path = output_dir / "Evidence_Safety.md"
    change_path.write_text(
        "# Change Log\n\n"
        f"## Application archetype\n\n{change['archetype']}\n\n"
        f"## Formatting source\n\n{data['document_settings'].get('reference_template_name') or 'Career Centre smart default'}\n\n"
        "## Evidence included\n" + bullets(change["included_evidence"]) + "\n\n"
        "## Evidence excluded\n" + bullets(change["excluded_evidence"]) + "\n\n"
        "## Keywords deliberately included\n" + bullets(change["keywords_included"]) + "\n\n"
        "## Keywords omitted for evidence safety\n" + bullets(change["keywords_omitted"]) + "\n\n"
        "## Ambiguities\n" + bullets(change["ambiguities"]) + "\n",
        encoding="utf-8",
    )
    safety_lines = [
        "# Evidence Safety",
        "",
        "Every substantive tailored-document item was validated against the evidence IDs in Application_Pack_Input.json.",
        "",
        "## Evidence register",
    ]
    for evidence in data["evidence"]:
        safety_lines.append(
            f"- **{evidence['evidence_id']}** — {evidence['safe_wording']} "
            f"({evidence['source_type']}; {evidence['confidence']})"
        )
    safety_lines.extend(["", "## Unresolved ambiguity", bullets(change["ambiguities"])])
    safety_path.write_text("\n".join(safety_lines) + "\n", encoding="utf-8")
    return [change_path, safety_path]


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--cv-template")
    parser.add_argument("--cover-template")
    args = parser.parse_args()
    input_path = Path(args.input).resolve()
    data = json.loads(input_path.read_text(encoding="utf-8"))
    errors = validate_application_pack(data)
    if errors:
        print(json.dumps({"passed": False, "stage": "input_contract", "errors": errors}, indent=2))
        return 1

    skill_root = Path(__file__).resolve().parent.parent
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    template_name = data["document_settings"]["template"]
    default_cv = skill_root / "assets" / "templates" / (
        "Professional_ATS_CV_Template.docx" if template_name == "professional" else "Executive_ATS_CV_Template.docx"
    )
    default_cover = skill_root / "assets" / "templates" / "Cover_Letter_Template.docx"
    if template_name == "reference" and not args.cv_template:
        print(json.dumps({"passed": False, "stage": "template", "errors": ["Reference template mode requires --cv-template with a DOCX file."]}, indent=2))
        return 1
    cv_template = Path(args.cv_template).resolve() if args.cv_template else default_cv
    cover_template = Path(args.cover_template).resolve() if args.cover_template else default_cover
    if cv_template.suffix.casefold() != ".docx" or not cv_template.is_file():
        print(json.dumps({"passed": False, "stage": "template", "errors": ["CV template must be an existing DOCX file."]}, indent=2))
        return 1
    try:
        Document(cv_template)
    except Exception as exc:
        print(json.dumps({"passed": False, "stage": "template", "errors": [f"CV template could not be opened: {exc}"]}, indent=2))
        return 1

    preserved_input = output_dir / "Application_Pack_Input.json"
    shutil.copy2(input_path, preserved_input)
    cv = _build_cv(data, output_dir, cv_template)
    cover = _build_cover_letter(data, output_dir, cover_template)
    notes = _write_notes(data, output_dir)

    expected_breaks = data["document_settings"]["page_target"] - 1
    cv_rules = {
        "document_type": "cv",
        "minimum_font_pt": data["document_settings"]["minimum_font_pt"],
        "expected_manual_page_breaks": expected_breaks,
        "required_text": [section["heading"] for section in data["cv"]["sections"]],
    }
    cv_report = validate_docx(cv, cv_rules)
    cv_report_path = output_dir / "CV_Structural_Validation.json"
    cv_report_path.write_text(json.dumps(cv_report, indent=2), encoding="utf-8")
    reports = [cv_report_path]
    cover_report = None
    cover_report_path = None
    if cover:
        cover_report = validate_docx(
            cover,
            {
                "document_type": "cover_letter",
                "minimum_font_pt": data["document_settings"]["minimum_font_pt"],
                "required_text": [data["role"]["identity"]["company"], data["role"]["identity"]["title"]],
            },
        )
        cover_report_path = output_dir / "Cover_Letter_Structural_Validation.json"
        cover_report_path.write_text(json.dumps(cover_report, indent=2), encoding="utf-8")
        reports.append(cover_report_path)

    structural_pass = cv_report["passed"] and (cover_report is None or cover_report["passed"])
    if not data["document_settings"]["cv_only"] and not cover:
        structural_pass = False
    files: list[Path] = [preserved_input, cv, *notes, *reports]
    if cover:
        files.append(cover)
    manifest = {
        "schema_version": "4.0",
        "role_id": data["role"]["role_id"],
        "cv_only": data["document_settings"]["cv_only"],
        "status": "BUILT_STRUCTURAL_PASS" if structural_pass else "FAILED_STRUCTURAL_QA",
        "visual_qa": "pending",
        "files": [
            {
                "name": path.name,
                "path": str(path.relative_to(output_dir)),
                "sha256": _sha256(path),
            }
            for path in sorted(files)
        ],
        "cv": cv.name,
        "cover_letter": cover.name if cover else None,
        "structural_reports": [path.name for path in reports],
        "render_reports": [],
        "format_source": data["document_settings"].get("reference_template_name") or "smart_default",
    }
    manifest_path = output_dir / "application_pack_manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps(manifest, indent=2))
    return 0 if structural_pass else 1


if __name__ == "__main__":
    raise SystemExit(main())
