#!/usr/bin/env python3
"""Produce a transparent, non-scored CV diagnostic from text or DOCX."""
from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from pathlib import Path
from typing import Any


EMAIL_RE = re.compile(r"(?i)\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b")
PHONE_RE = re.compile(r"(?<!\w)(?:\+?\d[\d\s().-]{7,}\d)(?!\w)")
URL_RE = re.compile(r"(?i)\b(?:https?://|www\.)\S+")
BULLET_RE = re.compile(r"^\s*(?:[•●▪◦*-]|\d+[.)])\s+")
DATE_RE = re.compile(
    r"(?i)\b(?:19|20)\d{2}\b|\b(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|"
    r"may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|"
    r"dec(?:ember)?)\s+(?:19|20)\d{2}\b"
)
OUTCOME_RE = re.compile(
    r"(?i)(?:\b\d+(?:\.\d+)?\s?%|[$£€¥₹]\s?\d|\b(?:USD|AUD|CAD|GBP|EUR|NZD|SGD|INR)\s?\d|"
    r"\b\d+[xX]\b|\b\d{2,}\+?\s+(?:people|employees|users|customers|sites|markets|countries|projects)\b)"
)
PASSIVE_RE = re.compile(r"(?i)\b(?:was|were|is|are|been|being)\s+(?:\w+ly\s+)?\w+(?:ed|en)\b")
GENERIC_PHRASES = (
    "results-driven",
    "results oriented",
    "dynamic professional",
    "proven track record",
    "excellent communication skills",
    "hard-working",
    "hardworking",
    "team player",
    "strategic thinker",
    "detail-oriented",
    "fast-paced environment",
    "responsible for",
    "various stakeholders",
    "cutting-edge",
    "best-in-class",
    "leveraged synergies",
)
HEADING_GROUPS = {
    "summary": ("summary", "professional summary", "profile", "career profile", "executive profile"),
    "experience": ("experience", "professional experience", "work experience", "career history", "employment history"),
    "skills": ("skills", "core skills", "key skills", "capabilities", "areas of expertise", "competencies"),
    "education": ("education", "qualifications", "academic background"),
    "projects": ("projects", "selected projects", "portfolio"),
    "certifications": ("certifications", "certificates", "recognition and certification", "awards and certifications"),
}


def _normalise_line(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _heading_key(value: str) -> str:
    return re.sub(r"[^\w\s]", "", _normalise_line(value).casefold()).strip()


def _docx_paragraph_text(paragraph: Any) -> str:
    text = paragraph.text
    if not text.strip() or BULLET_RE.match(text):
        return text
    style_name = str(getattr(getattr(paragraph, "style", None), "name", "")).casefold()
    properties = getattr(paragraph._p, "pPr", None)
    has_numbering = properties is not None and getattr(properties, "numPr", None) is not None
    if has_numbering or "bullet" in style_name or "list" in style_name:
        return "• " + text.strip()
    return text


def _finding(
    category: str,
    impact: str,
    message: str,
    recommendation: str,
    evidence: str | None = None,
) -> dict[str, Any]:
    item: dict[str, Any] = {
        "category": category,
        "impact": impact,
        "message": message,
        "recommendation": recommendation,
    }
    if evidence:
        item["evidence"] = evidence[:240]
    return item


def extract_text(path: Path) -> str:
    """Read supported source text without sending it to an external service."""
    suffix = path.suffix.casefold()
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - environment-specific recovery
            raise RuntimeError("python-docx is required to review DOCX files") from exc
        document = Document(path)
        blocks = [_docx_paragraph_text(paragraph) for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(blocks)
    if suffix in {".txt", ".md", ".text"}:
        return path.read_text(encoding="utf-8")
    raise ValueError("Supported review inputs are DOCX, TXT and Markdown. Extract PDF text first.")


def review_text(text: str, *, source: str = "supplied text", language: str = "en") -> dict[str, Any]:
    """Return qualitative findings, strengths and observable metrics—never an ATS score."""
    lines = [_normalise_line(line) for line in text.splitlines() if _normalise_line(line)]
    words = re.findall(r"\b[^\W_]+(?:['’.-][^\W_]+)*\b", text, flags=re.UNICODE)
    word_count = len(words)
    english_rules = language.casefold().startswith("en")
    heading_lookup = {
        alias: group
        for group, aliases in HEADING_GROUPS.items()
        for alias in aliases
    }
    headings_found = sorted(
        {
            heading_lookup[key]
            for line in lines
            if (key := _heading_key(line)) in heading_lookup
        }
    ) if english_rules else []
    bullets = [BULLET_RE.sub("", line).strip() for line in lines if BULLET_RE.match(line)]
    duplicate_counts = Counter(
        re.sub(r"[^\w\s]", "", bullet.casefold()).strip()
        for bullet in bullets
        if len(bullet.split()) >= 5
    )
    duplicate_bullets = sorted(key for key, count in duplicate_counts.items() if key and count > 1)
    date_signals = len(DATE_RE.findall(text))
    outcome_signals = len(OUTCOME_RE.findall(text))
    emails = sorted(set(EMAIL_RE.findall(text)))
    phones = sorted(set(PHONE_RE.findall(text)))
    raw_urls = sorted(set(URL_RE.findall(text)))
    sentences = [
        sentence.strip()
        for sentence in re.split(r"(?<=[.!?])\s+|\n+", text)
        if len(sentence.strip().split()) >= 8
    ]
    long_sentences = [
        sentence for sentence in sentences
        if len(sentence.split()) > 38 and sentence.count("|") < 3
    ]
    generic_hits = [phrase for phrase in GENERIC_PHRASES if phrase in text.casefold()] if english_rules else []
    passive_hits = PASSIVE_RE.findall(text) if english_rules else []
    starter_counts = Counter(
        re.match(r"[A-Za-zÀ-ÖØ-öø-ÿ]+", bullet).group(0).casefold()
        for bullet in bullets
        if re.match(r"[A-Za-zÀ-ÖØ-öø-ÿ]+", bullet)
    )
    repeated_starters = sorted(word for word, count in starter_counts.items() if count >= 3)

    findings: list[dict[str, Any]] = []
    strengths: list[str] = []

    if word_count < 80:
        findings.append(_finding(
            "extraction",
            "high",
            "There is too little selectable text for a reliable CV review.",
            "Check that the file is readable and not an image-only scan, then provide the full CV text.",
        ))
    elif word_count < 180:
        findings.append(_finding(
            "content depth",
            "medium",
            "The extracted CV is unusually brief, so important evidence may be missing.",
            "Check whether roles, dates, achievements or qualifications were lost during extraction.",
        ))
    else:
        strengths.append(f"The file yielded {word_count} words of selectable text for review.")

    if emails:
        strengths.append("An email address is available in selectable text.")
    else:
        findings.append(_finding(
            "contact details",
            "high",
            "No email address was detected in selectable text.",
            "Add a current email address as real text rather than placing it only in an image or inaccessible header.",
        ))
    if not phones and word_count >= 180:
        findings.append(_finding(
            "contact details",
            "low",
            "No phone number was detected in selectable text.",
            "Include one if it is appropriate for the target market, or keep it intentionally hidden as a confirmed preference.",
        ))

    if english_rules:
        if len(headings_found) >= 3:
            strengths.append("Common CV sections are clearly detectable in the extracted text.")
        elif word_count >= 180:
            findings.append(_finding(
                "structure",
                "high" if len(headings_found) < 2 else "medium",
                "Few standard section headings were detected, which may make the CV harder to scan or parse.",
                "Use short, conventional headings for summary, experience, skills and education where those sections apply.",
                ", ".join(headings_found) if headings_found else "No common English section headings detected",
            ))
    elif word_count >= 80:
        strengths.append("The text was preserved without forcing English-only writing rules.")

    if date_signals >= 2:
        strengths.append("Multiple chronology signals are visible in selectable text.")
    elif word_count >= 180:
        findings.append(_finding(
            "chronology",
            "medium",
            "The review found few visible date signals.",
            "Make employment and education dates easy to associate with the correct role or qualification.",
        ))

    if len(bullets) >= 4:
        strengths.append("The CV uses scannable bullet-style evidence rather than only dense paragraphs.")
    if outcome_signals >= 2:
        strengths.append("The CV includes observable scale or outcome signals; these should still be checked against source evidence.")
    if duplicate_bullets:
        findings.append(_finding(
            "repetition",
            "medium",
            f"{len(duplicate_bullets)} bullet statement(s) appear more than once.",
            "Keep the strongest version once and use the recovered space for distinct evidence.",
            duplicate_bullets[0],
        ))
    if repeated_starters:
        findings.append(_finding(
            "writing variety",
            "low",
            "Several bullets begin with the same word, which can make the evidence feel repetitive.",
            "Vary the opening only where it improves accuracy; do not replace precise verbs with inflated synonyms.",
            ", ".join(repeated_starters[:5]),
        ))
    if long_sentences:
        findings.append(_finding(
            "readability",
            "medium",
            f"{len(long_sentences)} sentence(s) exceed 38 words.",
            "Split the longest sentences so the action, context and result can be understood in one scan.",
            long_sentences[0],
        ))
    if generic_hits:
        findings.append(_finding(
            "specificity",
            "medium" if len(generic_hits) >= 2 else "low",
            "Generic career phrasing may be taking space from distinctive evidence.",
            "Replace only the flagged phrases with specific scope, decisions, stakeholders, outputs or outcomes that the source supports.",
            ", ".join(generic_hits[:6]),
        ))
    if len(passive_hits) >= 3:
        findings.append(_finding(
            "ownership clarity",
            "low",
            "Several passive constructions may make ownership less clear.",
            "Review them manually and use active wording only where the evidence supports personal ownership.",
            f"{len(passive_hits)} possible passive constructions",
        ))
    if raw_urls:
        findings.append(_finding(
            "links",
            "low",
            "Visible raw web addresses were detected.",
            "In the final Word CV, use descriptive linked text where practical and confirm every link target.",
            raw_urls[0],
        ))

    impact_order = {"high": 0, "medium": 1, "low": 2}
    findings.sort(key=lambda item: (impact_order[item["impact"]], item["category"]))
    if word_count < 80:
        status = "insufficient_text"
    elif any(item["impact"] in {"high", "medium"} for item in findings):
        status = "needs_attention"
    else:
        status = "clear"
    return {
        "source": source,
        "review_type": "qualitative_cv_diagnostic",
        "no_universal_ats_score": True,
        "status": status,
        "language": language,
        "strengths": strengths[:6],
        "findings": findings,
        "metrics": {
            "word_count": word_count,
            "nonempty_lines": len(lines),
            "bullet_count": len(bullets),
            "headings_found": headings_found,
            "date_signals": date_signals,
            "outcome_or_scale_signals": outcome_signals,
            "email_addresses_detected": len(emails),
            "phone_numbers_detected": len(phones),
            "visible_raw_urls": len(raw_urls),
        },
        "limitations": [
            "This is a transparent writing and parseability diagnostic, not an employer ATS score.",
            "Every finding needs career-context and evidence review before editing.",
        ],
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="DOCX, TXT or Markdown CV")
    parser.add_argument("--language", default="en", help="Document language code; English rules run only for en*")
    parser.add_argument("--json-out", type=Path)
    args = parser.parse_args()
    try:
        text = extract_text(args.input)
        report = review_text(text, source=str(args.input), language=args.language)
    except Exception as exc:
        print(json.dumps({"status": "error", "error": str(exc)}, indent=2, ensure_ascii=False))
        return 2
    rendered = json.dumps(report, indent=2, ensure_ascii=False)
    if args.json_out:
        args.json_out.write_text(rendered + "\n", encoding="utf-8")
    print(rendered)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
